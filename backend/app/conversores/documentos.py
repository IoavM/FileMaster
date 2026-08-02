"""
Conversor de documentos.
Soporta conversiones entre: PDF, DOCX, XLSX, TXT, CSV.
Usa PyMuPDF (fitz), python-docx, openpyxl y utilidades estándar.
"""
from io import BytesIO, StringIO
import csv
import zipfile
import docx
import openpyxl
import fitz

def pdf_a_txt(pdf_bytes: bytes) -> bytes:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    texto = ""
    for pagina in doc:
        texto += pagina.get_text()
    return texto.encode("utf-8")

def pdf_a_docx(pdf_bytes: bytes) -> bytes:
    # Extraer texto de PDF y escribirlo en un documento Word (.docx)
    doc_fitz = fitz.open(stream=pdf_bytes, filetype="pdf")
    doc_docx = docx.Document()
    
    for pagina in doc_fitz:
        lineas = pagina.get_text().splitlines()
        for linea in lineas:
            if linea.strip():
                doc_docx.add_paragraph(linea)
                
    buffer = BytesIO()
    doc_docx.save(buffer)
    return buffer.getvalue()

def pdf_a_imagen(pdf_bytes: bytes, formato: str = "png") -> bytes:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    fmt = formato.lower()
    fmt_fitz = "jpeg" if fmt in ("jpg", "jpeg") else fmt
    matriz = fitz.Matrix(2.0, 2.0)  # Renderizado de alta definición (200 DPI)
    
    if len(doc) == 1:
        pagina = doc[0]
        pix = pagina.get_pixmap(matrix=matriz)
        return pix.tobytes(fmt_fitz)
    else:
        # Si tiene más de una página, las empaquetamos en un ZIP
        buffer_zip = BytesIO()
        with zipfile.ZipFile(buffer_zip, "w", zipfile.ZIP_DEFLATED) as archivo_zip:
            for i, pagina in enumerate(doc):
                pix = pagina.get_pixmap(matrix=matriz)
                archivo_zip.writestr(f"pagina_{i+1}.{fmt}", pix.tobytes(fmt_fitz))
        return buffer_zip.getvalue()

def docx_a_pdf(docx_bytes: bytes) -> bytes:
    # Convertir DOCX a PDF dibujando texto básico usando PyMuPDF
    doc = docx.Document(BytesIO(docx_bytes))
    pdf = fitz.open()
    pagina = pdf.new_page()
    
    x = 50
    y = 50
    alto_pagina = pagina.rect.height
    
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            if y > alto_pagina - 50:
                pagina = pdf.new_page()
                y = 50
            pagina.insert_text((x, y), parrafo.text, fontsize=11)
            y += 20
            
    buffer = BytesIO()
    pdf.save(buffer)
    return buffer.getvalue()

def docx_a_txt(docx_bytes: bytes) -> bytes:
    doc = docx.Document(BytesIO(docx_bytes))
    texto = "\n".join([p.text for p in doc.paragraphs])
    return texto.encode("utf-8")

def txt_a_pdf(txt_bytes: bytes) -> bytes:
    contenido = txt_bytes.decode("utf-8", errors="ignore")
    pdf = fitz.open()
    pagina = pdf.new_page()
    
    x = 50
    y = 50
    alto_pagina = pagina.rect.height
    
    for linea in contenido.splitlines():
        if y > alto_pagina - 50:
            pagina = pdf.new_page()
            y = 50
        pagina.insert_text((x, y), linea, fontsize=10)
        y += 18
        
    buffer = BytesIO()
    pdf.save(buffer)
    return buffer.getvalue()

def txt_a_docx(txt_bytes: bytes) -> bytes:
    contenido = txt_bytes.decode("utf-8", errors="ignore")
    doc = docx.Document()
    for linea in contenido.splitlines():
        doc.add_paragraph(linea)
        
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def xlsx_a_csv(xlsx_bytes: bytes) -> bytes:
    wb = openpyxl.load_workbook(BytesIO(xlsx_bytes), data_only=True)
    hoja = wb.active
    
    buffer_str = StringIO()
    escritor = csv.writer(buffer_str)
    
    for fila in hoja.iter_rows(values_only=True):
        fila_limpia = ["" if v is None else str(v) for v in fila]
        escritor.writerow(fila_limpia)
        
    return buffer_str.getvalue().encode("utf-8")

def xlsx_a_pdf(xlsx_bytes: bytes) -> bytes:
    # Convertir XLSX a PDF tabulado simple
    wb = openpyxl.load_workbook(BytesIO(xlsx_bytes), data_only=True)
    hoja = wb.active
    
    pdf = fitz.open()
    pagina = pdf.new_page()
    
    x = 50
    y = 50
    alto_pagina = pagina.rect.height
    
    for fila in hoja.iter_rows(values_only=True):
        if y > alto_pagina - 50:
            pagina = pdf.new_page()
            y = 50
            
        texto_fila = " | ".join(["" if v is None else str(v) for v in fila[:6]]) # Mostrar máx 6 columnas
        pagina.insert_text((x, y), texto_fila, fontsize=9)
        y += 20
        
    buffer = BytesIO()
    pdf.save(buffer)
    return buffer.getvalue()

def csv_a_xlsx(csv_bytes: bytes) -> bytes:
    contenido = csv_bytes.decode("utf-8", errors="ignore")
    lector = csv.reader(StringIO(contenido))
    
    wb = openpyxl.Workbook()
    hoja = wb.active
    
    for r_idx, fila in enumerate(lector, start=1):
        for c_idx, valor in enumerate(fila, start=1):
            # Intentar convertir tipos
            try:
                if "." in valor:
                    valor = float(valor)
                else:
                    valor = int(valor)
            except ValueError:
                pass
            hoja.cell(row=r_idx, column=c_idx, value=valor)
            
    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()

def csv_a_pdf(csv_bytes: bytes) -> bytes:
    contenido = csv_bytes.decode("utf-8", errors="ignore")
    lector = csv.reader(StringIO(contenido))
    
    pdf = fitz.open()
    pagina = pdf.new_page()
    
    x = 50
    y = 50
    alto_pagina = pagina.rect.height
    
    for fila in lector:
        if y > alto_pagina - 50:
            pagina = pdf.new_page()
            y = 50
        texto_fila = " | ".join(fila[:6])
        pagina.insert_text((x, y), texto_fila, fontsize=9)
        y += 20
        
    buffer = BytesIO()
    pdf.save(buffer)
    return buffer.getvalue()

def convertir_documento(bytes_entrada: bytes, ext_entrada: str, ext_salida: str) -> bytes:
    """
    Función despachadora para conversiones de documentos.
    """
    ext_entrada = ext_entrada.lower()
    ext_salida = ext_salida.lower()
    
    if ext_entrada == "pdf":
        if ext_salida == "txt":
            return pdf_a_txt(bytes_entrada)
        elif ext_salida == "docx":
            return pdf_a_docx(bytes_entrada)
        elif ext_salida in ("png", "jpg", "jpeg"):
            return pdf_a_imagen(bytes_entrada, ext_salida)
            
    elif ext_entrada == "docx":
        if ext_salida == "pdf":
            return docx_a_pdf(bytes_entrada)
        elif ext_salida == "txt":
            return docx_a_txt(bytes_entrada)
            
    elif ext_entrada == "txt":
        if ext_salida == "pdf":
            return txt_a_pdf(bytes_entrada)
        elif ext_salida == "docx":
            return txt_a_docx(bytes_entrada)
            
    elif ext_entrada == "xlsx":
        if ext_salida == "csv":
            return xlsx_a_csv(bytes_entrada)
        elif ext_salida == "pdf":
            return xlsx_a_pdf(bytes_entrada)
            
    elif ext_entrada == "csv":
        if ext_salida == "xlsx":
            return csv_a_xlsx(bytes_entrada)
        elif ext_salida == "pdf":
            return csv_a_pdf(bytes_entrada)
            
    raise ValueError(f"Conversión de documento no soportada de {ext_entrada} a {ext_salida}")
